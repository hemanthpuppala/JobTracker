import json
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DEFAULT_MARGIN = 720  # twips
DEFAULT_PAGE_WIDTH = 12240
DEFAULT_FONT = "Calibri"
DEFAULT_BASE_FONT_SIZE = 10  # pt


def _resolve_style(style_config: dict | None) -> dict:
    """Resolve style config with defaults. Sizes in half-points for OOXML."""
    cfg = style_config or {}
    font = cfg.get("font", DEFAULT_FONT)
    base = cfg.get("base_font_size", DEFAULT_BASE_FONT_SIZE)
    ratio = base / DEFAULT_BASE_FONT_SIZE

    # Margins: frontend sends pt, convert to inches for python-docx
    margins = cfg.get("margins")
    margin_inches = 0.5  # default
    if margins:
        margin_inches = margins.get("left", 36) / 72  # pt to inches

    return {
        "font": font,
        "name_size": round(14 * ratio, 1),
        "contact_size": round(9 * ratio, 1),
        "header_size": round(10 * ratio, 1),
        "title_size": round(10 * ratio, 1),
        "body_size": round(9.5 * ratio, 1),
        "margin_inches": margin_inches,
    }


def _add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_right_tab(paragraph, position=DEFAULT_PAGE_WIDTH - DEFAULT_MARGIN * 2):
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(position))
    tabs.append(tab)
    pPr.append(tabs)


def _set_spacing(paragraph, before=0, after=0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before * 0.5)
    pf.space_after = Pt(after * 0.5)


def _add_run(paragraph, text, bold=False, italic=False, size=19, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size * 0.5)
    run.font.name = font
    return run


def _parse_bullets(raw) -> list[str]:
    """Parse bullets from JSON string or list."""
    if isinstance(raw, str):
        return json.loads(raw)
    if isinstance(raw, list):
        return [b if isinstance(b, str) else b.get("text", "") for b in raw]
    return []


def _add_bullet_list(doc, bullets: list[str], st: dict):
    """Add a list of bullet paragraphs with consistent styling."""
    for b in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        _set_spacing(bp, 0, 0)
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bp.runs:
            bp.runs[0].text = b
            bp.runs[0].font.size = Pt(st["body_size"])
            bp.runs[0].font.name = st["font"]
        else:
            _add_run(bp, b, size=st["body_size"] * 2, font=st["font"])


def _add_section_header(doc, title: str, st: dict):
    """Add a section header with bottom border."""
    p = doc.add_paragraph()
    _set_spacing(p, 80, 40)
    _add_bottom_border(p)
    _add_run(p, title, bold=True, size=st["header_size"] * 2, font=st["font"])


def generate_resume(
    profile: dict,
    experiences: list[dict],
    projects: list[dict],
    skills: list[dict],
    education: list[dict],
    custom_summary: str | None = None,
    style_config: dict | None = None,
) -> bytes:
    st = _resolve_style(style_config)
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Emu(DEFAULT_PAGE_WIDTH * 914)
    section.page_height = Emu(15840 * 914)
    section.top_margin = Inches(st["margin_inches"])
    section.bottom_margin = Inches(st["margin_inches"])
    section.left_margin = Inches(st["margin_inches"])
    section.right_margin = Inches(st["margin_inches"])

    # Content width for right-aligned tabs
    margin_twips = int(st["margin_inches"] * 1440)
    content_width = DEFAULT_PAGE_WIDTH - margin_twips * 2

    # Set default font
    style = doc.styles['Normal']
    style.font.name = st["font"]
    style.font.size = Pt(st["body_size"])

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, 0, 0)
    _add_run(p, profile.get("full_name", ""), bold=True, size=st["name_size"] * 2, font=st["font"])

    # Contact line
    contact_parts = [
        profile.get("location"), profile.get("phone"), profile.get("email"),
        profile.get("linkedin"), profile.get("github"), profile.get("portfolio"),
    ]
    contact_line = "  |  ".join(part for part in contact_parts if part)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, 40, 60)
    _add_run(p, contact_line, size=st["contact_size"] * 2, font=st["font"])

    # Summary
    summary_text = custom_summary or profile.get("summary", "")
    if summary_text:
        _add_section_header(doc, "SUMMARY", st)
        p = doc.add_paragraph()
        _set_spacing(p, 40, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_run(p, summary_text, size=st["body_size"] * 2, font=st["font"])

    # Technical Skills
    if skills:
        _add_section_header(doc, "TECHNICAL SKILLS", st)
        for skill in sorted(skills, key=lambda s: s.get("sort_order", 0)):
            p = doc.add_paragraph()
            _set_spacing(p, 0, 0)
            _add_run(p, skill["category"] + ": ", bold=True, size=st["body_size"] * 2, font=st["font"])
            _add_run(p, skill["items"], size=st["body_size"] * 2, font=st["font"])

    # Professional Experience
    if experiences:
        _add_section_header(doc, "PROFESSIONAL EXPERIENCE", st)
        for exp in sorted(experiences, key=lambda e: e.get("sort_order", 0)):
            date_end = exp.get("date_end") or "Present"
            date_range = f"{exp['date_start']} to {date_end}"

            p = doc.add_paragraph()
            _set_spacing(p, 60, 0)
            _add_right_tab(p, content_width)
            _add_run(p, f"{exp['title']} | {exp['company']}", bold=True, size=st["title_size"] * 2, font=st["font"])
            _add_run(p, f"\t{date_range}", size=st["title_size"] * 2, font=st["font"])

            _add_bullet_list(doc, _parse_bullets(exp["bullets"]), st)

    # Projects
    if projects:
        _add_section_header(doc, "PROJECTS", st)
        for proj in sorted(projects, key=lambda p: p.get("sort_order", 0)):
            p = doc.add_paragraph()
            _set_spacing(p, 60, 0)
            _add_run(p, proj["name"], bold=True, size=st["title_size"] * 2, font=st["font"])
            _add_run(p, f" | {proj['tech_stack']}", italic=True, size=st["body_size"] * 2, font=st["font"])

            _add_bullet_list(doc, _parse_bullets(proj["bullets"]), st)

    # Education
    if education:
        _add_section_header(doc, "EDUCATION", st)
        for edu in sorted(education, key=lambda e: e.get("sort_order", 0)):
            date_range = f"{edu['date_start']} to {edu['date_end']}"
            gpa_str = f" (GPA: {edu['gpa']})" if edu.get("gpa") else ""

            p = doc.add_paragraph()
            _set_spacing(p, 40, 0)
            _add_right_tab(p, content_width)
            _add_run(p, f"{edu['degree']}{gpa_str}  |  {edu['institution']}", bold=True, size=st["body_size"] * 2, font=st["font"])
            _add_run(p, f"\t{date_range}", size=st["body_size"] * 2, font=st["font"])

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
